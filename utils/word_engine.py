import uuid
import logging
import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.text.paragraph import Paragraph
from docx.table import Table

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class WordFragmentEngine:
    """
    WordFragmentEngine handles granular Word document manipulation, 
    focusing on UUID-based traceability and style-safe assembly.
    """
    
    def __init__(self, master_path=None):
        """
        Initialize the engine with an optional master template.
        """
        self.master_doc = Document(master_path) if master_path else Document()
        self._synced_styles = set() # Cache of style IDs already processed

    def _get_style_id(self, style_element):
        return style_element.get(qn('w:styleId'))

    def inject_uuid_to_paragraph(self, paragraph, fragment_id=None):
        """
        Inject a Word Bookmark as a UUID anchor into a paragraph.
        This provides traceability without altering visible text.
        """
        if not fragment_id:
            fragment_id = f"FRAG_{uuid.uuid4().hex[:12]}"
        
        p = paragraph._p
        # Word expects a unique integer ID for bookmarks within a document
        # We use a hash of the UUID string for consistency
        bm_id = str(abs(hash(fragment_id)) % (10**9))
        
        # Create <w:bookmarkStart>
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), bm_id)
        start.set(qn('w:name'), fragment_id)
        
        # Create <w:bookmarkEnd>
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), bm_id)
        
        # Inject at the beginning of the paragraph
        # Ideally, we should wrap the content, but for "anchoring", start/end at beginning is safer for extraction.
        p.insert(0, start)
        p.append(end)
        
        return fragment_id

    def iter_block_items(self, parent):
        """
        Iterates over child elements of a parent (like Document or Table Cell) 
        in correct visual order.
        """
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        
        parent_elm = parent.element.body if hasattr(parent.element, 'body') else parent.element

        for child in parent_elm.iterchildren():
            if child.tag == qn('w:p'):
                yield Paragraph(child, parent)
            elif child.tag == qn('w:tbl'):
                yield Table(child, parent)

    def extract_fragments_by_heading(self, source_path):
        """
        Extract document fragments delimited by Heading styles.
        Supports both Paragraphs and Tables.
        """
        doc = Document(source_path)
        fragments = []
        current_fragment = None

        for item in self.iter_block_items(doc):
            # Detect Heading levels (only for paragraphs)
            if hasattr(item, 'style') and item.style.name.startswith('Heading'):
                if current_fragment:
                    fragments.append(current_fragment)
                
                fid = self.inject_uuid_to_paragraph(item)
                current_fragment = {
                    "id": fid,
                    "title": item.text,
                    "style": item.style.name,
                    "content_elements": [item]
                }
            else:
                if current_fragment:
                    current_fragment["content_elements"].append(item)
                else:
                    # Prelude fragment
                    fid = self.inject_uuid_to_paragraph(item) if hasattr(item, 'text') else f"FRAG_{uuid.uuid4().hex[:12]}"
                    current_fragment = {
                        "id": fid,
                        "title": "Prelude",
                        "style": "Normal",
                        "content_elements": [item]
                    }
        
        if current_fragment:
            fragments.append(current_fragment)
            
        return fragments

    def _get_style_element(self, doc, name_or_id):
        """Finds the <w:style> element for a given style name or style ID."""
        for style in doc.styles.element.xpath('w:style'):
            # Check for Style ID attribute
            if style.get(qn('w:styleId')) == name_or_id:
                return style
            # Check for Name element
            name_el = style.find(qn('w:name'))
            if name_el is not None and name_el.get(qn('w:val')) == name_or_id:
                return style
        return None

    def sync_style_deep(self, source_doc, source_name_or_id, target_doc):
        """
        [Advanced] Deeply clones the XML definition of a style from source to target,
        including its parent styles (recursive dependency handling with cycle detection).
        """
        # 1. Find source XML to get the canonical Style ID
        source_style_el = self._get_style_element(source_doc, source_name_or_id)
        if source_style_el is None:
            return target_doc.styles['Normal']
        
        style_id = self._get_style_id(source_style_el)

        # 2. Check if already processed in this session or exists in target
        if style_id in self._synced_styles:
            return target_doc.styles[style_id]
        
        # Check target XML directly for the ID to avoid python-docx name mapping issues
        for style in target_doc.styles.element.xpath('w:style'):
            if self._get_style_id(style) == style_id:
                self._synced_styles.add(style_id)
                return target_doc.styles[style_id]

        # 3. Cycle Detection & Dependency Handling
        self._synced_styles.add(style_id) # Mark as visiting
        
        based_on = source_style_el.find(qn('w:basedOn'))
        if based_on is not None:
            parent_id = based_on.get(qn('w:val'))
            if parent_id and parent_id != style_id:
                self.sync_style_deep(source_doc, parent_id, target_doc)

        # 4. Clone and Append
        new_style_el = copy.deepcopy(source_style_el)
        target_doc.styles.element.append(new_style_el)
        
        logger.info(f"Deep synced style: {style_id}")
        return target_doc.styles[style_id]

    def stitch_fragment(self, source_doc, fragment, target_doc=None):
        """
        Appends a fragment (Paragraphs and Tables) to the target document.
        """
        if target_doc is None:
            target_doc = self.master_doc
            
        for element in fragment["content_elements"]:
            if isinstance(element, Table):
                # Handle Table stitching
                new_tbl = target_doc.add_table(rows=len(element.rows), cols=len(element.columns))
                # Sync table style (Try ID first, then name)
                try:
                    target_style = self.sync_style_deep(source_doc, element.style.name, target_doc)
                    new_tbl.style = target_style
                except:
                    pass
                
                # Copy cell contents (simplified - for complex tables we'd need nested iteration)
                for r_idx, row in enumerate(element.rows):
                    for c_idx, cell in enumerate(row.cells):
                        new_tbl.cell(r_idx, c_idx).text = cell.text
            else:
                # Handle Paragraph stitching
                new_p = target_doc.add_paragraph(element.text)
                try:
                    new_p.style = self.sync_style_deep(source_doc, element.style.name, target_doc)
                except Exception as e:
                    logger.error(f"Style sync failed: {e}")
                    new_p.style = target_doc.styles['Normal']
                
                self.inject_uuid_to_paragraph(new_p, fragment["id"])

    def save(self, path):
        self.master_doc.save(path)
