import os
import uuid
import logging
import numpy as np
from docx import Document
from utils.word_engine import WordFragmentEngine

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DynamicSemanticChunker:
    """
    动态语义分段引擎：
    结合 Word 物理结构（Heading）与 语义内容（Embedding Similarity）进行智能切片。
    确保每个 Fragment 既符合文档架构，又具有完整的语义上下文。
    """
    def __init__(self, embedder=None, threshold=0.85, target_dir="data/chunks"):
        self.embedder = embedder # 应为具备 embed_documents 方法的对象
        self.threshold = threshold
        self.target_dir = target_dir
        self.word_engine = WordFragmentEngine()
        os.makedirs(self.target_dir, exist_ok=True)

    def _calculate_similarity(self, vec1, vec2):
        if vec1 is None or vec2 is None: return 1.0
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        if norm1 == 0 or norm2 == 0: return 0.0
        return np.dot(vec1, vec2) / (norm1 * norm2)

    def chunk_document(self, docx_path):
        """
        核心分段逻辑：
        1. 提取所有段落及其文本。
        2. 计算相邻段落相似度。
        3. 综合 Heading 标记与相似度突变点进行切分。
        """
        doc = Document(docx_path)
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        
        if not paragraphs:
            return []

        # 预计算所有段落的 Embedding (批量处理以优化性能)
        embeddings = []
        if self.embedder:
            try:
                embeddings = self.embedder.embed_documents([p.text for p in paragraphs])
            except Exception as e:
                logger.error(f"Embedding calculation failed: {e}")
                embeddings = [None] * len(paragraphs)
        else:
            embeddings = [None] * len(paragraphs)

        chunks = []
        current_chunk_paras = [paragraphs[0]]
        
        for i in range(1, len(paragraphs)):
            prev_p = paragraphs[i-1]
            curr_p = paragraphs[i]
            
            # 策略 A: 强制 Heading 分段 (物理边界)
            is_heading_split = curr_p.style.name.startswith('Heading')
            
            # 策略 B: 语义相似度分段 (内容边界)
            similarity = self._calculate_similarity(embeddings[i-1], embeddings[i])
            is_semantic_split = similarity < self.threshold
            
            if is_heading_split or is_semantic_split:
                # 结束当前 Chunk，开启新 Chunk
                chunks.append(self._create_chunk_file(current_chunk_paras))
                current_chunk_paras = [curr_p]
            else:
                current_chunk_paras.append(curr_p)
                
        # 补全最后一个 Chunk
        if current_chunk_paras:
            chunks.append(self._create_chunk_file(current_chunk_paras))
            
        return chunks

    def _create_chunk_file(self, paragraphs):
        """
        利用 WordFragmentEngine 保持样式与注入 UUID。
        """
        chunk_id = f"SEM_{uuid.uuid4().hex[:12]}"
        fname = f"{chunk_id}.docx"
        fpath = os.path.join(self.target_dir, fname)
        
        new_doc = Document()
        for p in paragraphs:
            # 注入书签 UUID
            self.word_engine.inject_uuid_to_paragraph(p, chunk_id)
            # 添加到新文档并同步样式 (简化版同步)
            new_p = new_doc.add_paragraph(p.text)
            new_p.style = p.style.name if p.style.name in [s.name for s in new_doc.styles] else 'Normal'
            
        new_doc.save(fpath)
        logger.info(f"Created semantic chunk: {fname} (Size: {len(paragraphs)} paras)")
        return {"id": chunk_id, "path": fpath, "text": "\n".join([p.text for p in paragraphs])}

class MockEmbedder:
    """用于测试的 Mock Embedder"""
    def embed_documents(self, texts):
        # 返回随机向量或占位符
        return [np.random.rand(1536) for _ in texts]

if __name__ == "__main__":
    # 快速验证
    chunker = DynamicSemanticChunker(embedder=MockEmbedder())
    print("SemanticChunker Engine Ready.")
