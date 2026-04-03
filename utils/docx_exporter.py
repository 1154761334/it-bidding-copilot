"""
Word 文档导出工具
使用 python-docx 将投标内容填入 Word 模板
"""
import io
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def create_bid_document(
    project_name: str,
    company_name: str,
    sections: dict,
    template_path: Optional[str] = None,
) -> io.BytesIO:
    """
    生成投标书 Word 文档

    Args:
        project_name: 项目名称
        company_name: 投标企业名称
        sections: 章节内容 {section_key: {title, chapters: {ch_key: {title, content}}}}
        template_path: 可选的 Word 模板路径

    Returns:
        BytesIO 包含 .docx 文件内容
    """
    if not HAS_DOCX:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    if template_path:
        doc = Document(template_path)
    else:
        doc = Document()
        _setup_styles(doc)

    # ── 封面 ──
    _add_cover_page(doc, project_name, company_name)

    # ── 目录 ──
    doc.add_heading("目  录", level=1)
    doc.add_paragraph("（请使用 Word 目录功能自动生成）")
    doc.add_page_break()

    # ── 正文 ──
    for sec_key, sec in sections.items():
        doc.add_heading(sec["title"], level=1)
        for ch_key, ch in sec.get("chapters", {}).items():
            doc.add_heading(ch["title"], level=2)
            _add_content_paragraphs(doc, ch["content"])
        doc.add_page_break()

    # ── 附件清单 ──
    doc.add_heading("附件清单", level=1)
    appendices = [
        "附件一：企业营业执照扫描件",
        "附件二：资质证书扫描件",
        "附件三：近三年审计财务报告",
        "附件四：类似项目业绩合同复印件",
        "附件五：法人授权书",
    ]
    for a in appendices:
        doc.add_paragraph(a, style="List Number")

    # ── 输出 ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _setup_styles(doc: "Document"):
    """配置文档默认样式"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    # 标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f"Heading {i}"]
        heading_style.font.name = "黑体"
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1D, 0x23)


def _add_cover_page(doc: "Document", project_name: str, company_name: str):
    """生成封面页"""
    # 空行
    for _ in range(6):
        doc.add_paragraph("")

    # 项目名称
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(project_name)
    run.font.size = Pt(22)
    run.bold = True

    # 投标文件
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("投  标  文  件")
    run.font.size = Pt(28)
    run.bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    # 企业信息
    info_items = [
        f"投标单位：{company_name}",
        f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}",
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(16)

    doc.add_page_break()


def _add_content_paragraphs(doc: "Document", content: str):
    """将文本内容按段落添加到文档"""
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 检测表格（简单的 Markdown 表格识别）
        if stripped.startswith("|") and stripped.endswith("|"):
            # 跳过分隔行
            if set(stripped.replace("|", "").replace("-", "").strip()) <= {" ", ""}:
                continue
            doc.add_paragraph(stripped, style="Normal")
        elif stripped.startswith("#"):
            # 子标题
            level = min(stripped.count("#"), 4)
            text = stripped.lstrip("# ")
            doc.add_heading(text, level=min(level + 1, 4))
        else:
            doc.add_paragraph(stripped)
