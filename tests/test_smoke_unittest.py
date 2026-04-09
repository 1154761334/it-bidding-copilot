import asyncio
import base64
import datetime
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from unittest.mock import AsyncMock, patch

import main
from api.routers import config_v2
from api.routers import dashboard_v2, drafting_v2
from api.routers import enterprise_v2, rfp_v2
from api.models.assets_v2 import Company
from api.models.assets_v2 import CompanyAsset, EnterpriseCase, EnterpriseCertificate, EnterprisePersonnel, SourceDocument
from api.models.bid_draft_v2 import BidDraft
from api.models.rfp_v2 import RFPProject, RFPRequirement
from api.services.model_runtime_service import get_model_runtime_info
from api.services.enterprise_ingest_service import EnterpriseIngestService
from api.services.drafting_workflow import DraftingWorkflow
from api.services.context_service import (
    get_latest_draft_for_project,
    get_latest_project,
    get_or_create_primary_company,
    get_primary_company,
)
from api.services.bid_exporter import BidExporter
from api.services import drafting_task_service
from api.services.enterprise_asset_service import EnterpriseAssetService
from api.services.drafting_review_service import DraftingReviewService
from api.services.task_registry import InMemoryTaskRegistry
from utils.rfp_analyzer import RFPAnalyzer


class QueryStub:
    def __init__(self, result):
        self.result = result
        self.filter_called = False

    def order_by(self, *_args, **_kwargs):
        return self

    def filter(self, *_args, **_kwargs):
        self.filter_called = True
        return self

    def first(self):
        return self.result

    def get(self, _id):
        return self.result

    def scalar(self):
        return self.result

    def all(self):
        return self.result if isinstance(self.result, list) else [self.result]

    def limit(self, *_args, **_kwargs):
        return self

    def count(self):
        return len(self.result) if isinstance(self.result, list) else (1 if self.result else 0)

    def distinct(self):
        return self


class SessionStub:
    def __init__(self, result=None):
        self.result = result
        self.added = None
        self.committed = False
        self.refreshed = None
        self.last_query = None

    def query(self, _model):
        self.last_query = QueryStub(self.result)
        return self.last_query

    def add(self, value):
        self.added = value

    def commit(self):
        self.committed = True

    def refresh(self, value):
        self.refreshed = value


class AppSmokeTests(unittest.TestCase):
    def test_main_app_imports(self):
        self.assertEqual(main.app.title, "IT Bidding Copilot Industrial API")

    def test_root_handler_returns_operational_status(self):
        response = asyncio.run(main.root())

        self.assertEqual(response["status"], "operational")


class DashboardRouteTests(unittest.TestCase):
    def test_dashboard_context_handler_returns_latest_context(self):
        company = Company(id=11, company_name="测试企业")
        project = RFPProject(id=22, project_name="测试项目")
        draft = BidDraft(id=33)

        class DashboardSessionStub:
            def close(self):
                return None

        with patch("api.routers.dashboard_v2.SessionLocal", return_value=DashboardSessionStub()), \
            patch("api.routers.dashboard_v2.get_primary_company", return_value=company), \
            patch("api.routers.dashboard_v2.get_latest_project", return_value=project), \
            patch("api.routers.dashboard_v2.get_latest_draft_for_project", return_value=draft):
            response = asyncio.run(dashboard_v2.get_dashboard_context())

        self.assertEqual(
            response,
            {
                "current_company_id": 11,
                "current_company_name": "测试企业",
                "current_project_id": 22,
                "current_project_name": "测试项目",
                "current_draft_id": "33",
            },
        )


class ContextServiceTests(unittest.TestCase):
    def test_get_primary_company_returns_first_company(self):
        company = Company(company_name="测试企业")
        session = SessionStub(company)

        result = get_primary_company(session)

        self.assertIs(result, company)

    def test_get_or_create_primary_company_creates_when_missing(self):
        session = SessionStub(None)

        result = get_or_create_primary_company(session)

        self.assertIsInstance(result, Company)
        self.assertEqual(result.company_name, "未命名企业")
        self.assertIs(session.added, result)
        self.assertTrue(session.committed)
        self.assertIs(session.refreshed, result)

    def test_get_latest_project_filters_by_company(self):
        project = RFPProject(project_name="项目A")
        session = SessionStub(project)

        result = get_latest_project(session, company_id=7)

        self.assertIs(result, project)
        self.assertTrue(session.last_query.filter_called)

    def test_get_latest_draft_for_project_returns_first_draft(self):
        draft = BidDraft(section_title="第一章")
        session = SessionStub(draft)

        result = get_latest_draft_for_project(session, project_id=3)

        self.assertIs(result, draft)
        self.assertTrue(session.last_query.filter_called)


class ReviewRouteQueryStub:
    def __init__(self, drafts):
        self.drafts = drafts

    def filter(self, *_args, **_kwargs):
        return self

    def all(self):
        return self.drafts


class ReviewRouteSessionStub:
    def __init__(self, drafts):
        self.drafts = drafts

    def query(self, _model):
        return ReviewRouteQueryStub(self.drafts)


class ReviewRouteTests(unittest.TestCase):
    def test_review_handler_returns_structured_summary(self):
        drafts = [
            BidDraft(
                id=1,
                project_id=9,
                section_title="服务方案",
                generation_status="COMPLETED",
                content_markdown="## 服务方案\n投标人提供完整响应。",
                audit_logs={"final_feedback": "APPROVED"},
                source_fragments=["源片段 A"],
                winning_points="优势 A",
            ),
            BidDraft(
                id=2,
                project_id=9,
                section_title="SLA 条款",
                generation_status="COMPLETED",
                content_markdown="## SLA\n当前版本缺少赔付条款。",
                audit_logs={"final_feedback": "FIX_REQUIRED: 缺少赔付条款"},
                source_fragments=["源片段 B"],
                winning_points="优势 B",
            ),
        ]
        payload = asyncio.run(DraftingReviewService(ReviewRouteSessionStub(drafts)).run_red_team_review(9))
        self.assertEqual(payload["project_id"], 9)
        self.assertEqual(payload["total_drafts"], 2)
        self.assertEqual(payload["approved_drafts"], 1)
        self.assertEqual(len(payload["section_reviews"]), 2)
        self.assertEqual(payload["section_reviews"][0]["verdict"], "APPROVED")
        self.assertEqual(payload["section_reviews"][1]["verdict"], "REJECTED")
        self.assertEqual(payload["section_reviews"][1]["source_fragments"], ["源片段 B"])

    def test_review_handler_marks_incomplete_draft_as_rejected(self):
        drafts = [
            BidDraft(
                id=1,
                project_id=9,
                section_title="服务方案",
                generation_status="REVIEWING",
                content_markdown="",
                audit_logs={"final_feedback": ""},
                source_fragments=[],
            ),
        ]
        payload = asyncio.run(DraftingReviewService(ReviewRouteSessionStub(drafts)).run_red_team_review(9))
        self.assertEqual(payload["approved_drafts"], 0)
        self.assertEqual(payload["section_reviews"][0]["verdict"], "REJECTED")
        self.assertIn("禁止作为最终标书导出", payload["section_reviews"][0]["feedback"])

    def test_export_readiness_marks_incomplete_sections(self):
        project = RFPProject(id=9, project_name="测试项目", status="DEVIATION_CONFIRMED")
        drafts = [
            BidDraft(
                id=1,
                project_id=9,
                section_title="服务方案",
                generation_status="COMPLETED",
                content_markdown="正文",
                source_fragments=["证据A"],
            ),
            BidDraft(
                id=2,
                project_id=9,
                section_title="商务条款",
                generation_status="REVIEWING",
                content_markdown="",
                source_fragments=[],
            ),
        ]
        payload = drafting_v2.build_export_readiness(project, drafts)
        self.assertFalse(payload["ready"])
        self.assertEqual(payload["rejected_sections"][0]["section_title"], "商务条款")
        self.assertFalse(next(item for item in payload["checks"] if item["key"] == "all_drafts_completed")["passed"])

    def test_export_readiness_reports_template_and_image_evidence(self):
        project = RFPProject(id=10, project_name="测试项目", status="COMPLETED")
        drafts = [
            BidDraft(
                id=1,
                project_id=10,
                section_title="技术方案",
                generation_status="COMPLETED",
                content_markdown="正文\n[IMAGE:/tmp/proof.png]",
                source_fragments=["[IMAGE:/tmp/proof.png]"],
                audit_logs={"final_feedback": "APPROVED"},
            ),
        ]
        payload = DraftingReviewService(None).build_export_readiness(project, drafts, master_template_available=True)
        template_check = next(item for item in payload["checks"] if item["key"] == "master_template_available")
        image_check = next(item for item in payload["checks"] if item["key"] == "image_evidence_ready")
        self.assertTrue(template_check["passed"])
        self.assertTrue(image_check["passed"])
        self.assertEqual(image_check["detail"]["image_evidence_count"], 2)

    def test_update_draft_content_increments_version_and_sets_reviewing(self):
        draft = BidDraft(
            id=15,
            project_id=9,
            section_title="实施方案",
            generation_status="PENDING",
            content_markdown="",
            version=1,
            last_updated=datetime.date(2026, 4, 1),
        )

        class DraftQueryStub:
            def filter(self, *_args, **_kwargs):
                return self

            def first(self):
                return draft

        class DraftSessionStub:
            def __init__(self):
                self.committed = False
                self.refreshed = None

            def query(self, _model):
                return DraftQueryStub()

            def commit(self):
                self.committed = True

            def refresh(self, value):
                self.refreshed = value

        session = DraftSessionStub()
        payload = asyncio.run(
            drafting_v2.update_draft_content(
                15,
                drafting_v2.DraftContentUpdateRequest(content_markdown="## 实施方案\n已补齐现场服务内容。"),
                session,
            )
        )

        self.assertEqual(draft.content_markdown, "## 实施方案\n已补齐现场服务内容。")
        self.assertEqual(draft.version, 2)
        self.assertEqual(draft.generation_status, "REVIEWING")
        self.assertTrue(session.committed)
        self.assertEqual(payload["status"], "saved")
        self.assertEqual(payload["version"], 2)


class EnterpriseRouteOverviewTests(unittest.TestCase):
    def test_assets_overview_builder_returns_counts_and_items(self):
        company = Company(id=7, company_name="测试企业")

        class CountQueryStub:
            def __init__(self, items):
                self.items = items

            def filter(self, *_args, **_kwargs):
                return self

            def order_by(self, *_args, **_kwargs):
                return self

            def limit(self, value):
                self.items = self.items[:value]
                return self

            def all(self):
                return self.items

            def count(self):
                return len(self.items)

        service = EnterpriseAssetService(SessionStub(None))
        payload = service.build_assets_overview(company)
        self.assertEqual(payload["company_id"], 7)
        self.assertEqual(payload["counts"]["certificates"], 1)
        self.assertEqual(payload["counts"]["cases"], 1)
        self.assertEqual(payload["counts"]["personnel"], 1)
        self.assertEqual(payload["counts"]["source_documents"], 1)
        self.assertEqual(payload["counts"]["images"], 1)
        self.assertEqual(payload["certificates"][0]["raw_name"], "ISO9001")
        self.assertEqual(payload["cases"][0]["project_name"], "政务云项目")

    def test_enterprise_intake_readiness_builder_returns_checklist(self):
        company = Company(
            id=7,
            company_name="测试企业",
            unified_social_credit_code="91330100TEST00001X",
        )

        class CountQueryStub:
            def __init__(self, count_value):
                self.count_value = count_value

            def filter(self, *_args, **_kwargs):
                return self

            def count(self):
                return self.count_value

        service = EnterpriseAssetService(SessionStub(None))
        payload = service.build_enterprise_intake_readiness(company)
        self.assertTrue(payload["ready"])
        self.assertEqual(payload["company_name"], "测试企业")
        self.assertEqual(len(payload["checks"]), 6)

    def test_latest_ingest_batch_builder_returns_latest_non_rfp_batch(self):
        company = Company(id=7, company_name="测试企业")
        latest_docs = [
            SourceDocument(id=21, company_id=7, filename="商务技术文件.docx", file_type="BUSINESS_DOC", local_path="/tmp/business.docx", upload_date=datetime.date(2026, 4, 9)),
            SourceDocument(id=22, company_id=7, filename="政务云案例.md", file_type="CASE", local_path="/tmp/case.md", upload_date=datetime.date(2026, 4, 9)),
        ]

        class UploadDateQueryStub:
            def filter(self, *_args, **_kwargs):
                return self

            def order_by(self, *_args, **_kwargs):
                return self

            def limit(self, *_args, **_kwargs):
                return self

            def scalar(self):
                return datetime.date(2026, 4, 9)

        class CountQueryStub:
            def __init__(self, count_value):
                self.count_value = count_value

            def filter(self, *_args, **_kwargs):
                return self

            def count(self):
                return self.count_value

        service = EnterpriseAssetService(SessionStub(None))
        payload = service.build_latest_ingest_batch(company)
        self.assertTrue(payload["has_batch"])
        self.assertEqual(payload["batch_date"], "2026-04-09")
        self.assertEqual(payload["counts"]["source_documents"], 2)
        self.assertEqual(payload["counts"]["certificates"], 3)
        self.assertEqual(payload["counts"]["cases"], 2)
        self.assertEqual(payload["counts"]["images"], 8)

    def test_assets_browser_builder_filters_by_kind_and_query(self):
        company = Company(id=7, company_name="测试企业")

        class QueryStub:
            def __init__(self, items):
                self.items = items

            def filter(self, *_args, **_kwargs):
                return self

            def order_by(self, *_args, **_kwargs):
                return self

            def all(self):
                return self.items

        service = EnterpriseAssetService(SessionStub())
        payload = service.build_assets_browser(company, asset_kind="certificate", query="ISO")
        self.assertEqual(payload["total"], 1)
        self.assertEqual(payload["items"][0]["kind"], "certificate")
        self.assertEqual(payload["items"][0]["title"], "ISO9001")


class TaskRegistryTests(unittest.TestCase):
    def test_registry_create_update_and_get(self):
        async def scenario():
            registry = InMemoryTaskRegistry()
            await registry.create("task-1", stage="queued")
            await registry.update("task-1", status="running", stage="parsing_document")
            return await registry.get("task-1")

        record = asyncio.run(scenario())
        self.assertIsNotNone(record)
        self.assertEqual(record.task_id, "task-1")
        self.assertEqual(record.status, "running")
        self.assertEqual(record.stage, "parsing_document")


class RfpRouteTests(unittest.TestCase):
    def test_task_status_returns_registry_state(self):
        async def scenario():
            with patch.object(rfp_v2.task_registry, "get", return_value=type("TaskRecordStub", (), {
                "status": "running",
                "stage": "parsing_document",
                "result": None,
                "error": None,
                "to_dict": lambda self: {
                    "task_id": "rfp_demo_task",
                    "status": "running",
                    "stage": "parsing_document",
                    "created_at": "2026-04-08T00:00:00+00:00",
                    "updated_at": "2026-04-08T00:00:01+00:00",
                },
            })()):
                return await rfp_v2.get_task_status("rfp_demo_task", SessionStub(None))

        payload = asyncio.run(scenario())
        self.assertEqual(payload["status"], "running")
        self.assertEqual(payload["stage"], "parsing_document")
        self.assertIn("created_at", payload)
        self.assertIn("updated_at", payload)

    def test_task_status_returns_failed_for_unknown_task(self):
        async def scenario():
            with patch.object(rfp_v2.task_registry, "get", return_value=None):
                return await rfp_v2.get_task_status("rfp_missing_task", SessionStub(None))

        payload = asyncio.run(scenario())
        self.assertEqual(payload["status"], "failed")
        self.assertEqual(payload["error"], "Task not found")

    def test_analysis_check_returns_quality_report(self):
        project = RFPProject(id=3, project_name="测试项目", budget=1200000)
        requirement = RFPRequirement(
            project_id=3,
            original_section="第五章 评标办法",
            clause_index="5.1",
            category="TECHNICAL",
            description="提供技术响应方案",
            is_fatal=True,
            max_score=10,
            evidence_required="方案说明",
        )

        class QueryStub:
            def __init__(self, result):
                self.result = result

            def filter(self, *_args, **_kwargs):
                return self

            def first(self):
                return self.result if not isinstance(self.result, list) else self.result[0]

            def all(self):
                return self.result if isinstance(self.result, list) else [self.result]

        class AnalysisCheckSessionStub:
            def query(self, model):
                if model is RFPProject:
                    return QueryStub(project)
                return QueryStub([requirement])

        payload = asyncio.run(rfp_v2.get_analysis_check(3, AnalysisCheckSessionStub()))
        self.assertEqual(payload["project_id"], 3)
        self.assertIn("quality_report", payload)
        self.assertIn("checks", payload["quality_report"])

    def test_update_deviation_matrix_updates_comments_and_status(self):
        project = RFPProject(id=5, project_name="项目A")
        requirement = RFPRequirement(id=10, project_id=5, description="要求A", match_comment="旧备注", match_status="FAIL")

        class QueryStub:
            def __init__(self, items):
                self.items = items

            def filter(self, *_args, **_kwargs):
                return self

            def first(self):
                return self.items[0] if self.items else None

            def all(self):
                return self.items

        class DeviationSessionStub:
            def __init__(self):
                self.committed = False

            def query(self, model):
                if model is RFPProject:
                    return QueryStub([project])
                return QueryStub([requirement])

            def commit(self):
                self.committed = True

        payload = asyncio.run(
            rfp_v2.update_deviation_matrix(
                5,
                rfp_v2.DeviationUpdateRequest(items=[rfp_v2.DeviationItemUpdate(id=10, resp="人工修订应答", status="compliant")]),
                DeviationSessionStub(),
            )
        )
        self.assertEqual(payload["status"], "success")
        self.assertEqual(requirement.match_comment, "人工修订应答")
        self.assertEqual(requirement.match_status, "PASS")

    def test_confirm_deviation_matrix_updates_project_status(self):
        project = RFPProject(id=6, project_name="项目B", status="MATCHED")
        requirements = [
            RFPRequirement(id=1, project_id=6, description="要求1", match_comment="已填写"),
            RFPRequirement(id=2, project_id=6, description="要求2", match_comment="已填写"),
        ]

        class QueryStub:
            def __init__(self, *, first_result=None, all_result=None):
                self.first_result = first_result
                self.all_result = all_result or []

            def filter(self, *_args, **_kwargs):
                return self

            def first(self):
                return self.first_result

            def all(self):
                return self.all_result

        class ConfirmSessionStub:
            def __init__(self):
                self.committed = False

            def query(self, model):
                if model is RFPProject:
                    return QueryStub(first_result=project)
                return QueryStub(all_result=requirements)

            def commit(self):
                self.committed = True

        payload = asyncio.run(rfp_v2.confirm_deviation_matrix(6, ConfirmSessionStub()))
        self.assertEqual(payload["status"], "success")
        self.assertEqual(payload["project_status"], "DEVIATION_CONFIRMED")


class DraftingRouteTests(unittest.TestCase):
    def test_start_project_drafting_returns_pending(self):
        class CountQueryStub:
            def filter(self, *_args, **_kwargs):
                return self

            def count(self):
                return 2

        class DraftSessionStub:
            def query(self, _model):
                return CountQueryStub()

        async def scenario():
            with patch("api.routers.drafting_v2.start_project_draft_generation", new=AsyncMock(return_value="draft_project_demo")):
                return await drafting_v2.start_project_drafting(
                    7,
                    drafting_v2.ProjectDraftBatchRequest(max_sections=3),
                    DraftSessionStub(),
                )

        payload = asyncio.run(scenario())
        self.assertEqual(payload["status"], "pending")
        self.assertEqual(payload["task_id"], "draft_project_demo")
        self.assertEqual(payload["max_sections"], 3)
        self.assertFalse(payload["only_incomplete"])

    def test_start_project_drafting_supports_only_incomplete_mode(self):
        class CountQueryStub:
            def filter(self, *_args, **_kwargs):
                return self

            def count(self):
                return 4

        class DraftSessionStub:
            def query(self, _model):
                return CountQueryStub()

        async def scenario():
            with patch("api.routers.drafting_v2.start_project_draft_generation", new=AsyncMock(return_value="draft_project_retry")) as mocked:
                payload = await drafting_v2.start_project_drafting(
                    9,
                    drafting_v2.ProjectDraftBatchRequest(only_incomplete=True),
                    DraftSessionStub(),
                )
                return payload, mocked.await_args.kwargs

        payload, kwargs = asyncio.run(scenario())
        self.assertEqual(payload["task_id"], "draft_project_retry")
        self.assertTrue(payload["only_incomplete"])
        self.assertTrue(kwargs["only_incomplete"])

    def test_start_project_drafting_raises_when_no_drafts_exist(self):
        class CountQueryStub:
            def filter(self, *_args, **_kwargs):
                return self

            def count(self):
                return 0

        class DraftSessionStub:
            def query(self, _model):
                return CountQueryStub()

        async def scenario():
            return await drafting_v2.start_project_drafting(7, None, DraftSessionStub())

        with self.assertRaises(Exception):
            asyncio.run(scenario())

    def test_draft_task_status_returns_registry_state(self):
        async def scenario():
            with patch.object(drafting_v2.task_registry, "get", return_value=type("TaskRecordStub", (), {
                "status": "running",
                "stage": "researching",
                "result": None,
                "error": None,
                "to_dict": lambda self: {
                    "task_id": "draft_demo_task",
                    "status": "running",
                    "stage": "researching",
                    "created_at": "2026-04-08T00:00:00+00:00",
                    "updated_at": "2026-04-08T00:00:01+00:00",
                },
            })()):
                return await drafting_v2.get_draft_task_status("draft_demo_task")

        payload = asyncio.run(scenario())
        self.assertEqual(payload["status"], "running")
        self.assertEqual(payload["stage"], "researching")
        self.assertIn("created_at", payload)
        self.assertIn("updated_at", payload)

    def test_draft_task_status_returns_failed_for_unknown_task(self):
        async def scenario():
            with patch.object(drafting_v2.task_registry, "get", return_value=None):
                return await drafting_v2.get_draft_task_status("draft_missing_task")

        payload = asyncio.run(scenario())
        self.assertEqual(payload["status"], "failed")
        self.assertEqual(payload["error"], "Task not found")


class ConfigRouteTests(unittest.TestCase):
    def test_get_capabilities_returns_runtime_info(self):
        payload = asyncio.run(config_v2.get_model_capabilities())
        self.assertIn("provider", payload)
        self.assertIn("fallbacks", payload)
        self.assertIn("chat_enabled", payload)


class ModelRuntimeTests(unittest.TestCase):
    def test_runtime_info_marks_embedding_disabled_without_model(self):
        with patch("api.services.model_runtime_service.get_settings") as mock_get_settings:
            mock_get_settings.return_value = type(
                "SettingsStub",
                (),
                {
                    "LLM_PROVIDER": "openai-compatible",
                    "resolved_llm_api_key": "demo-key",
                    "resolved_llm_base_url": "https://ark.cn-beijing.volces.com/api/coding/v3",
                    "resolved_llm_model": "Auto",
                    "EMBEDDING_MODEL": "",
                },
            )()
            payload = get_model_runtime_info()

        self.assertTrue(payload["chat_enabled"])
        self.assertFalse(payload["embedding_enabled"])
        self.assertTrue(payload["fallbacks"]["embedding_zero_vector"])


class EnterpriseVaultTests(unittest.TestCase):
    def test_infer_asset_type_from_path_uses_folder_hints(self):
        service = EnterpriseIngestService(SessionStub(None))
        self.assertEqual(service._infer_asset_type_from_path(Path("/tmp/企业知识库/资质证书/ISO9001.md")), "CERTIFICATE")
        self.assertEqual(service._infer_asset_type_from_path(Path("/tmp/企业知识库/项目案例/政务云案例.md")), "CASE")
        self.assertEqual(service._infer_asset_type_from_path(Path("/tmp/企业知识库/人员简历/张三.md")), "PERSONNEL")

    def test_vault_ingest_route_returns_service_payload(self):
        async def scenario():
            with patch("api.routers.enterprise_v2.EnterpriseIngestService") as service_cls:
                service = service_cls.return_value
                service.ingest_vault_directory = AsyncMock(return_value={
                    "status": "vault processing finished",
                    "files_total": 3,
                    "results": [],
                })
                payload = enterprise_v2.VaultIngestRequest(vault_path="/tmp/demo_vault")
                return await enterprise_v2.vault_ingest(1, payload, SessionStub(None))

        result = asyncio.run(scenario())
        self.assertEqual(result["status"], "vault processing finished")
        self.assertEqual(result["files_total"], 3)


class EnterpriseFlowTests(unittest.TestCase):
    def test_noise_requirement_filter_skips_table_rows_and_procurement_boilerplate(self):
        analyzer = RFPAnalyzer()
        self.assertTrue(analyzer._is_noise_requirement("| 1 | 私有云建设 | 1 | 项 | 详见需求 |"))
        self.assertTrue(analyzer._is_noise_requirement("欢迎国内合格的供应商前来投标。"))
        self.assertTrue(analyzer._is_noise_requirement("联系人：张老师  联系电话：0571-12345678"))
        self.assertTrue(analyzer._is_noise_requirement("报名时间：2026年4月1日至2026年4月10日，每日上午9:00-11:30"))
        self.assertFalse(analyzer._is_noise_requirement("要求云平台具备多租户网络隔离能力。"))

    def test_extract_requirements_deduplicates_equivalent_clauses(self):
        analyzer = RFPAnalyzer()
        markdown = """
# 第四章 技术要求
1. 要求云平台具备多租户网络隔离能力。
2. 要求云平台具备多租户网络隔离能力。
联系人：张老师  联系电话：0571-12345678
"""
        items = analyzer._extract_requirements(markdown)

        self.assertEqual(len(items), 1)
        self.assertEqual(items[0]["description"], "要求云平台具备多租户网络隔离能力")

    def test_noise_requirement_filter_skips_background_narrative(self):
        analyzer = RFPAnalyzer()
        self.assertTrue(
            analyzer._is_noise_requirement(
                "目前，公司业务系统主要运行在公有云和杭钢云，将打造一朵混合云以提供稳定可靠的数字化基础设施。",
                current_section="项目建设背景",
            )
        )
        self.assertFalse(
            analyzer._is_noise_requirement(
                "要求云平台提供稳定可靠的数字化基础设施能力。",
                current_section="项目建设背景",
            )
        )

    def test_vault_ingest_collects_supported_files(self):
        class DummyDb:
            def query(self, _model):
                return type("QueryStub", (), {"filter": lambda self, *_a, **_k: self, "first": lambda self: Company(id=1, company_name="测试企业")})()

        service = EnterpriseIngestService(DummyDb())

        captured = []

        async def fake_ingest_local_file(**kwargs):
            captured.append(kwargs["display_name"])
            return {"filename": kwargs["display_name"], "status": "ingested", "asset_type": kwargs.get("asset_type_hint") or "GENERAL", "chunks_count": 1}

        with TemporaryDirectory() as tmp_dir:
            vault = Path(tmp_dir) / "vault"
            (vault / "资质证书").mkdir(parents=True)
            (vault / "项目案例").mkdir(parents=True)
            (vault / "资质证书" / "ISO9001.md").write_text("# ISO9001\n有效期至 2028-12-31", encoding="utf-8")
            (vault / "项目案例" / "案例.txt").write_text("政务云项目案例", encoding="utf-8")
            (vault / "忽略目录").mkdir()
            (vault / "忽略目录" / "image.gif").write_text("not supported", encoding="utf-8")

            with patch.object(service, "ingest_local_file", side_effect=fake_ingest_local_file):
                result = asyncio.run(service.ingest_vault_directory(1, str(vault), str(Path(tmp_dir) / "uploads")))

        self.assertEqual(result["files_total"], 2)
        self.assertEqual(sorted(captured), ["ISO9001.md", "案例.txt"])

    def test_business_doc_ingest_route_returns_service_payload(self):
        async def scenario():
            with patch("api.routers.enterprise_v2.EnterpriseIngestService") as service_cls:
                service = service_cls.return_value
                service.ingest_business_document = AsyncMock(return_value={
                    "status": "ingested",
                    "asset_type": "BUSINESS_DOC",
                    "certificates_created": 5,
                    "images_registered": 10,
                })
                payload = enterprise_v2.BusinessDocIngestRequest(
                    file_path="/root/it-bidding-copilot/docs/商务技术文件.docx",
                    display_name="商务技术文件.docx",
                )
                return await enterprise_v2.business_doc_ingest(1, payload, SessionStub(None))

        result = asyncio.run(scenario())
        self.assertEqual(result["status"], "ingested")
        self.assertEqual(result["asset_type"], "BUSINESS_DOC")
        self.assertEqual(result["certificates_created"], 5)

    def test_update_certificate_route_updates_fields(self):
        certificate = EnterpriseCertificate(id=7, raw_name="旧证书", cert_type="旧类型", cert_level="旧等级")

        class QueryStub:
            def filter(self, *_args, **_kwargs):
                return self

            def first(self):
                return certificate

        class SessionStub:
            def __init__(self):
                self.committed = False
                self.refreshed = None

            def query(self, _model):
                return QueryStub()

            def commit(self):
                self.committed = True

            def refresh(self, value):
                self.refreshed = value

        payload = asyncio.run(
            enterprise_v2.update_certificate(
                7,
                enterprise_v2.CertificateUpdateRequest(raw_name="新证书", cert_type="ISO", cert_level="A"),
                SessionStub(),
            )
        )
        self.assertEqual(payload["status"], "updated")
        self.assertEqual(certificate.raw_name, "新证书")
        self.assertEqual(certificate.cert_type, "ISO")
        self.assertEqual(certificate.cert_level, "A")

    def test_delete_case_route_removes_record(self):
        case = EnterpriseCase(id=9, project_name="测试案例")

        class QueryStub:
            def filter(self, *_args, **_kwargs):
                return self

            def first(self):
                return case

        class SessionStub:
            def __init__(self):
                self.deleted = None
                self.committed = False

            def query(self, _model):
                return QueryStub()

            def delete(self, value):
                self.deleted = value

            def commit(self):
                self.committed = True

        session = SessionStub()
        payload = asyncio.run(enterprise_v2.delete_case(9, session))
        self.assertEqual(payload["status"], "deleted")
        self.assertIs(session.deleted, case)

    def test_create_personnel_route_creates_record(self):
        company = Company(id=3, company_name="测试企业")

        class QueryStub:
            def order_by(self, *_args, **_kwargs):
                return self

            def first(self):
                return company

        class SessionStub:
            def __init__(self):
                self.added = None
                self.committed = False
                self.refreshed = None

            def query(self, _model):
                return QueryStub()

            def add(self, value):
                self.added = value

            def commit(self):
                self.committed = True

            def refresh(self, value):
                self.refreshed = value

        session = SessionStub()
        payload = asyncio.run(
            enterprise_v2.create_personnel(
                enterprise_v2.PersonnelUpdateRequest(name="李四", role="项目经理", level="高级", years_of_experience=8),
                session,
            )
        )
        self.assertEqual(payload["status"], "created")
        self.assertEqual(session.added.name, "李四")
        self.assertEqual(session.added.company_id, 3)

    def test_batch_delete_assets_deletes_supported_kinds(self):
        certificate = EnterpriseCertificate(id=1, raw_name="证书A")
        case = EnterpriseCase(id=2, project_name="案例B")

        class QueryStub:
            def __init__(self, model):
                self.model = model

            def filter(self, *_args, **_kwargs):
                return self

            def first(self):
                if self.model is EnterpriseCertificate:
                    return certificate
                if self.model is EnterpriseCase:
                    return case
                return None

        class SessionStub:
            def __init__(self):
                self.deleted = []
                self.committed = False

            def query(self, model):
                return QueryStub(model)

            def delete(self, value):
                self.deleted.append(value)

            def commit(self):
                self.committed = True

        session = SessionStub()
        payload = asyncio.run(
            enterprise_v2.batch_delete_assets(
                enterprise_v2.AssetBatchDeleteRequest(
                    items=[
                        enterprise_v2.AssetBatchDeleteItem(kind="certificate", id=1),
                        enterprise_v2.AssetBatchDeleteItem(kind="case", id=2),
                        enterprise_v2.AssetBatchDeleteItem(kind="image", id=99),
                    ]
                ),
                session,
            )
        )
        self.assertEqual(payload["count"], 2)
        self.assertEqual(len(session.deleted), 2)


class BusinessDocExtractorTests(unittest.TestCase):
    def test_split_sections_keeps_heading_without_body(self):
        from utils.business_doc_asset_extractor import BusinessDocAssetExtractor

        markdown = "# 8.2 公司各项资质证书\n# 8.2.1 ISO9001质量管理体系证书\n正文"
        extracted = BusinessDocAssetExtractor().extract(markdown)

        titles = [section["title"] for section in extracted["sections"]]
        self.assertIn("8.2 公司各项资质证书", titles)
        self.assertIn("8.2.1 ISO9001质量管理体系证书", titles)
        cert_names = [item["cert_name"] for item in extracted["certificates"]]
        self.assertTrue(any("ISO9001质量管理体系证书" in name for name in cert_names))

    def test_extract_cases_from_performance_section(self):
        from utils.business_doc_asset_extractor import BusinessDocAssetExtractor

        markdown = "# 9.4.1.2项目负责人业绩证明\n1、基于信创体系的“国资云”平台构建及创新技术研究项目\n我司于2021年6月16日与杭州钢铁集团有限公司签订 基于信创体系的“国资云”平台构建及创新技术研究项目 合同。"
        extracted = BusinessDocAssetExtractor().extract(markdown)

        self.assertTrue(extracted["cases"])
        self.assertIn("国资云", extracted["cases"][0]["project_name"])

    def test_extract_certificate_image_paths(self):
        from utils.business_doc_asset_extractor import BusinessDocAssetExtractor

        markdown = "# 5.1营业执照\n[IMAGE:/tmp/cert-1.png]"
        extracted = BusinessDocAssetExtractor().extract(markdown)

        self.assertEqual(extracted["certificates"][0]["image_paths"], ["/tmp/cert-1.png"])

    def test_extract_certificate_dates_and_level(self):
        from utils.business_doc_asset_extractor import BusinessDocAssetExtractor

        markdown = "# 8.2.1 ISO9001质量管理体系证书\n证书等级：先进级\n发证日期：2023年1月2日\n有效期至：2026年1月2日"
        extracted = BusinessDocAssetExtractor().extract(markdown)

        self.assertEqual(extracted["certificates"][0]["cert_level"], "先进级")
        self.assertEqual(extracted["certificates"][0]["issue_date"], "2023-01-02")
        self.assertEqual(extracted["certificates"][0]["expiry_date"], "2026-01-02")

    def test_extract_personnel_from_listed_certificate_section(self):
        from utils.business_doc_asset_extractor import BusinessDocAssetExtractor

        markdown = (
            "# 9.4.3.3云平台原厂实施团队成员资质证书\n"
            "1、刘松涛简历及技术资质证书\n"
            "2、赫中翔简历及技术资质证书\n"
            "3、许功成简历及技术资质证书"
        )
        extracted = BusinessDocAssetExtractor().extract(markdown)

        names = [item["name"] for item in extracted["personnel"]]
        self.assertIn("刘松涛", names)
        self.assertIn("赫中翔", names)
        self.assertIn("许功成", names)

    def test_extract_personnel_skips_feature_description_sections(self):
        from utils.business_doc_asset_extractor import BusinessDocAssetExtractor

        markdown = (
            "# 9.3.2.6.7资源编排\n"
            "资源编排服务是一款帮助云计算用户简化云资源管理和自动化部署运维的服务。"
        )
        extracted = BusinessDocAssetExtractor().extract(markdown)

        self.assertEqual(extracted["personnel"], [])

    def test_business_asset_llm_extractor_falls_back_when_llm_unavailable(self):
        from utils.business_asset_llm_extractor import BusinessAssetLLMExtractor

        extractor = BusinessAssetLLMExtractor.__new__(BusinessAssetLLMExtractor)
        sections = [{"title": "5.1营业执照", "content": ""}]
        fallback_assets = {
            "certificates": [{"title": "5.1营业执照", "cert_name": "营业执照", "cert_type": "营业执照", "scope": "", "evidence_text_preview": "", "image_paths": []}],
            "cases": [],
            "personnel": [],
            "authorizations": [],
            "social_security": [],
        }

        async def scenario():
            with patch("utils.business_asset_llm_extractor.LLMClient.is_configured", return_value=False):
                return await BusinessAssetLLMExtractor.standardize(
                    extractor,
                    sections=sections,
                    fallback_assets=fallback_assets,
                )

        result = asyncio.run(scenario())
        self.assertEqual(result["trace"]["mode"], "fallback")
        self.assertEqual(result["assets"]["certificates"][0]["cert_name"], "营业执照")


class EnterpriseBusinessAssetSanitizationTests(unittest.TestCase):
    def test_sanitize_business_doc_assets_filters_noise_and_dedupes(self):
        service = EnterpriseIngestService(SessionStub(None))
        extracted = {
            "certificates": [
                {"title": "5.1营业执照", "cert_name": "营业执照", "evidence_text_preview": "", "cert_type": "营业执照"},
                {"title": "9.3.2.6.7资源编排", "cert_name": "资源编排", "evidence_text_preview": "功能模块介绍", "cert_type": "证书材料"},
                {"title": "5.1营业执照", "cert_name": "营业执照", "evidence_text_preview": "", "cert_type": "营业执照"},
            ],
            "cases": [
                {"title": "历史案例", "project_name": "某政务云建设项目", "description": "云平台案例", "industry": "政企"},
                {"title": "历史案例", "project_name": "某政务云建设项目", "description": "云平台案例", "industry": "政企"},
            ],
            "personnel": [
                {"title": "9.4.3.3云平台原厂实施团队成员资质证书", "name": "刘松涛", "role": "工程师", "content_preview": "实施团队"},
                {"title": "9.3.2.6.7资源编排", "name": "资源编排", "role": "人员材料", "content_preview": "功能说明"},
                {"title": "9.4.3.3云平台原厂实施团队成员资质证书", "name": "刘松涛", "role": "工程师", "content_preview": "实施团队"},
            ],
            "authorizations": [],
            "social_security": [],
            "llm_trace": {},
        }

        sanitized = service._sanitize_business_doc_assets(extracted)

        self.assertEqual(len(sanitized["certificates"]), 1)
        self.assertEqual(sanitized["certificates"][0]["cert_name"], "营业执照")
        self.assertEqual(len(sanitized["cases"]), 1)
        self.assertEqual(len(sanitized["personnel"]), 1)
        self.assertEqual(sanitized["personnel"][0]["name"], "刘松涛")


class RfpAnalyzerTests(unittest.TestCase):
    def test_extract_scoring_items_from_markdown_table(self):
        from utils.rfp_analyzer import RFPAnalyzer

        analyzer = RFPAnalyzer.__new__(RFPAnalyzer)
        markdown = """
## 第五章 评标办法
| 序号 | 评分因素 | 评分细则 | 分值（分） |
| --- | --- | --- | --- |
| 一 、商务资信（20分） | 一 、商务资信（20分） | 一 、商务资信（20分） | 一 、商务资信（20分） |
| 1 | 投标人资信情况 | 提供ISO9001证书，每项得1分。提供相关证明材料。 | 4 |
| 2 | 投标人类似项目业绩 | 提供合同证明材料，每个案例得2分。 | 10 |
"""
        items = RFPAnalyzer._extract_scoring_items(analyzer, markdown)

        self.assertEqual(len(items), 2)
        self.assertEqual(items[0]["max_score"], 4.0)
        self.assertEqual(items[0]["category"], "COMMERCIAL")
        self.assertIn("证明材料", items[0]["evidence_required"])


class RfpRouteTests(unittest.TestCase):
    def test_resolve_rfp_source_path_falls_back_to_docs(self):
        with TemporaryDirectory() as tmp_dir:
            docs_dir = Path(tmp_dir) / "docs"
            docs_dir.mkdir(parents=True)
            target = docs_dir / "招标文件.docx"
            target.write_text("placeholder", encoding="utf-8")

            source = SourceDocument(filename="招标文件.docx", local_path="/missing/path/招标文件.docx")
            with patch.object(rfp_v2, "PROJECT_ROOT", Path(tmp_dir)):
                resolved = rfp_v2._resolve_rfp_source_path(source)

        self.assertEqual(resolved, str(target))

    def test_analysis_check_survives_missing_source_file(self):
        project = RFPProject(id=12, project_name="测试项目", budget=1000000, rfp_source_id=88)
        requirements = [
            RFPRequirement(
                project_id=12,
                original_section="技术要求",
                clause_index="1",
                category="TECHNICAL",
                description="要求云平台支持多租户隔离",
                is_fatal=False,
                max_score=5,
                evidence_required="证明材料",
            )
        ]
        source = SourceDocument(id=88, filename="missing.docx", local_path="/missing.docx")

        class QueryStub:
            def __init__(self, model):
                self.model = model

            def filter(self, *_args, **_kwargs):
                return self

            def first(self):
                if self.model is RFPProject:
                    return project
                if self.model is SourceDocument:
                    return source
                return None

            def all(self):
                if self.model is RFPRequirement:
                    return requirements
                return []

        class SessionStub:
            def query(self, model):
                return QueryStub(model)

        payload = asyncio.run(rfp_v2.get_analysis_check(12, SessionStub()))
        self.assertEqual(payload["project_id"], 12)
        self.assertIn("quality_report", payload)
        self.assertEqual(payload["quality_report"]["metrics"]["requirements_total"], 1)

    def test_confirm_project_analysis_updates_project_and_requirements(self):
        project = RFPProject(id=12, project_name="旧项目", budget=1000000, status="ANALYZING")
        requirement = RFPRequirement(
            id=8,
            project_id=12,
            original_section="技术要求",
            clause_index="1",
            category="TECHNICAL",
            description="旧要求",
            is_fatal=False,
            max_score=3,
            evidence_required="",
        )

        class QueryStub:
            def __init__(self, model):
                self.model = model

            def filter(self, *_args, **_kwargs):
                return self

            def first(self):
                if self.model is RFPProject:
                    return project
                return None

            def all(self):
                if self.model is RFPRequirement:
                    return [requirement]
                return []

        class SessionStub:
            def query(self, model):
                return QueryStub(model)

            def commit(self):
                return None

        payload = asyncio.run(
            rfp_v2.confirm_project_analysis(
                12,
                rfp_v2.AnalysisConfirmRequest(
                    project_info=rfp_v2.AnalysisProjectInfoUpdate(name="新项目", budget=2000000, deadline="2026-04-30"),
                    requirements=[
                        rfp_v2.AnalysisRequirementUpdate(
                            id=8,
                            description="修正后的要求",
                            category="COMMERCIAL",
                            is_fatal=True,
                            evidence_required="合同和截图",
                            max_score=8,
                        )
                    ],
                ),
                SessionStub(),
            )
        )

        self.assertEqual(payload["status"], "success")
        self.assertEqual(project.project_name, "新项目")
        self.assertEqual(project.budget, 2000000)
        self.assertEqual(project.status, "ANALYSIS_CONFIRMED")
        self.assertEqual(requirement.description, "修正后的要求")
        self.assertEqual(requirement.category, "COMMERCIAL")
        self.assertTrue(requirement.is_fatal)
        self.assertEqual(requirement.evidence_required, "合同和截图")
        self.assertEqual(requirement.max_score, 8)
        self.assertEqual(payload["result"]["project_status"], "ANALYSIS_CONFIRMED")


class DraftingWorkflowTests(unittest.TestCase):
    def test_should_continue_stops_after_three_iterations(self):
        workflow = DraftingWorkflow.__new__(DraftingWorkflow)
        decision = DraftingWorkflow.should_continue(workflow, {"is_approved": False, "iteration_count": 3})
        self.assertEqual(decision, "end")

    def test_search_personnel_evidence_includes_social_security_image(self):
        workflow = DraftingWorkflow.__new__(DraftingWorkflow)

        class QueryStub:
            def filter(self, *_args, **_kwargs):
                return self

            def order_by(self, *_args, **_kwargs):
                return self

            def limit(self, *_args, **_kwargs):
                return self

            def all(self):
                return [
                    EnterprisePersonnel(
                        id=1,
                        name="张三",
                        role="项目经理",
                        level="高级",
                        resume_text="具有私有云实施经验",
                        social_security_image_url="/tmp/social.png",
                    )
                ]

        class SessionStub:
            def query(self, _model):
                return QueryStub()

        workflow.db = SessionStub()
        results = DraftingWorkflow._search_personnel_evidence(
            workflow,
            company_id=1,
            section_title="项目负责人配置",
            requirements=["需提供项目负责人和社保证明"],
        )
        self.assertEqual(len(results), 1)
        self.assertIn("[IMAGE:/tmp/social.png]", results[0])


class BidExporterTests(unittest.TestCase):
    def test_exporter_uses_master_template_and_appends_evidence(self):
        png_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9WnXl1cAAAAASUVORK5CYII="
        )

        with TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            template_path = tmp_path / "template.docx"
            image_path = tmp_path / "evidence.png"
            output_dir = Path("/root/it-bidding-copilot/exports")
            output_dir.mkdir(parents=True, exist_ok=True)

            from docx import Document

            template_doc = Document()
            template_doc.add_paragraph("模板页眉正文")
            template_doc.save(template_path)
            image_path.write_bytes(png_bytes)

            project = RFPProject(id=99, project_name="测试项目", rfp_source_id=201)
            source = SourceDocument(id=201, local_path=str(template_path), filename="template.docx")
            draft = BidDraft(
                id=301,
                project_id=99,
                section_title="第一章 项目概述",
                section_index="1",
                generation_status="COMPLETED",
                content_markdown=f"## 概述\n正文内容\n[IMAGE:{image_path}]",
                source_fragments=["证据片段 A", "证据片段 B"],
            )

            class QueryStub:
                def __init__(self, model):
                    self.model = model

                def filter(self, *_args, **_kwargs):
                    return self

                def order_by(self, *_args, **_kwargs):
                    return self

                def first(self):
                    if self.model is RFPProject:
                        return project
                    if self.model is SourceDocument:
                        return source
                    return draft

                def all(self):
                    return [draft]

            class SessionStub:
                def query(self, model):
                    return QueryStub(model)

            exporter = BidExporter(SessionStub())
            output_path = exporter.export_project_bid(99)

            self.assertTrue(Path(output_path).exists())
            exported = Document(output_path)
            all_text = "\n".join(p.text for p in exported.paragraphs)
            self.assertIn("测试项目", all_text)
            self.assertIn("章节证据附录", all_text)
            self.assertIn("证据片段 A", all_text)
            self.assertIn("图片证据：evidence.png", all_text)

    def test_exporter_appends_images_from_mixed_evidence_fragments(self):
        png_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9WnXl1cAAAAASUVORK5CYII="
        )

        with TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            template_path = tmp_path / "template.docx"
            image_path = tmp_path / "evidence_mixed.png"

            from docx import Document

            template_doc = Document()
            template_doc.add_paragraph("模板页眉正文")
            template_doc.save(template_path)
            image_path.write_bytes(png_bytes)

            project = RFPProject(id=100, project_name="测试项目2", rfp_source_id=202)
            source = SourceDocument(id=202, local_path=str(template_path), filename="template.docx")
            draft = BidDraft(
                id=302,
                project_id=100,
                section_title="第二章 佐证材料",
                section_index="2",
                generation_status="COMPLETED",
                content_markdown="正文内容",
                source_fragments=[f"营业执照扫描件如下\n[IMAGE:{image_path}]"],
            )

            class QueryStub:
                def __init__(self, model):
                    self.model = model

                def filter(self, *_args, **_kwargs):
                    return self

                def order_by(self, *_args, **_kwargs):
                    return self

                def first(self):
                    if self.model is RFPProject:
                        return project
                    if self.model is SourceDocument:
                        return source
                    return draft

                def all(self):
                    return [draft]

            class SessionStub:
                def query(self, model):
                    return QueryStub(model)

            exporter = BidExporter(SessionStub())
            output_path = exporter.export_project_bid(100)

            exported = Document(output_path)
            all_text = "\n".join(p.text for p in exported.paragraphs)
            self.assertIn("营业执照扫描件如下", all_text)
            self.assertIn("图片证据：evidence_mixed.png", all_text)


class DraftingTaskServiceTests(unittest.TestCase):
    def test_select_project_drafts_filters_to_incomplete_or_empty(self):
        completed = SimpleNamespace(generation_status="COMPLETED", content_markdown="正文")
        reviewing = SimpleNamespace(generation_status="REVIEWING", content_markdown="待复核")
        empty_completed = SimpleNamespace(generation_status="COMPLETED", content_markdown="   ")
        pending = SimpleNamespace(generation_status="PENDING", content_markdown=None)

        selected = drafting_task_service._select_project_drafts(
            [completed, reviewing, empty_completed, pending],
            only_incomplete=True,
        )

        self.assertEqual(selected, [reviewing, empty_completed, pending])

    def test_select_project_drafts_applies_limit_after_filter(self):
        drafts = [
            SimpleNamespace(generation_status="REVIEWING", content_markdown="1"),
            SimpleNamespace(generation_status="PENDING", content_markdown=None),
            SimpleNamespace(generation_status="COMPLETED", content_markdown="  "),
        ]

        selected = drafting_task_service._select_project_drafts(
            drafts,
            only_incomplete=True,
            max_drafts=2,
        )

        self.assertEqual(selected, drafts[:2])


if __name__ == "__main__":
    unittest.main()
