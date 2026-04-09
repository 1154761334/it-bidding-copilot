from utils.word_engine import WordFragmentEngine
from docx import Document
import os

def create_test_source(path):
    doc = Document()
    doc.add_heading('Technical Specification', level=1)
    doc.add_paragraph('This is the first technical fragment with details about the server.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Component'
    table.cell(0, 1).text = 'Specification'
    table.cell(1, 0).text = 'CPU'
    table.cell(1, 1).text = '64 Cores'
    
    doc.add_heading('Security Requirements', level=1)
    doc.add_paragraph('This is the second fragment about encryption and MFA.')
    
    # Add a custom style paragraph
    p = doc.add_paragraph('Important notice in custom style.')
    if 'Intense Quote' in doc.styles:
        p.style = doc.styles['Intense Quote']
        
    doc.save(path)
    print(f"Created test source: {path}")

def run_test():
    source_file = "test_source.docx"
    master_file = "test_master.docx"
    output_file = "test_output.docx"
    
    # Setup
    create_test_source(source_file)
    Document().save(master_file) # Empty master
    
    engine = WordFragmentEngine(master_file)
    
    # 1. Extraction with UUID injection
    print("Extracting fragments...")
    source_doc = Document(source_file)
    fragments = engine.extract_fragments_by_heading(source_file)
    
    for f in fragments:
        print(f"Found Fragment: ID={f['id']}, Title={f['title']}")
        # 2. Stitching into Master
        engine.stitch_fragment(source_doc, f)
    
    # 3. Save and Verify
    engine.save(output_file)
    print(f"Saved merged document to: {output_file}")
    
    # Verify UUIDs in the final document
    final_doc = Document(output_file)
    bookmarks = []
    for p in final_doc.paragraphs:
        for bookmark in p._p.xpath('.//w:bookmarkStart'):
            bookmarks.append(bookmark.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name'))
    
    print(f"Found Bookmarks (UUIDs) in output: {bookmarks}")

    # Verify Tables in output
    print(f"Found Tables in output: {len(final_doc.tables)}")
    if len(final_doc.tables) > 0:
        print(f"First Table Row 2 Cell 2: {final_doc.tables[0].cell(1, 1).text}")
        if "64 Cores" in final_doc.tables[0].cell(1, 1).text:
            print("TEST PASSED: Table content preserved.")
        else:
            print("TEST FAILED: Table content corrupted.")
    else:
        print("TEST FAILED: Table missing.")

    if len(bookmarks) >= 2:
        print("TEST PASSED: UUID Traceability preserved.")
    else:
        print("TEST FAILED: UUIDs missing.")

if __name__ == "__main__":
    run_test()
