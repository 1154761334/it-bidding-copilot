from utils.word_chunker import DynamicSemanticChunker, MockEmbedder
from api.engines.tone_calibrator import ToneCalibrator
from docx import Document
import os

def create_integrated_test_source(path):
    doc = Document()
    doc.add_heading('System Architecture', level=1)
    doc.add_paragraph('总之，本系统非常有弹性，它很快，不仅如此还很稳。')
    doc.add_paragraph('这种方法能保护很多大数据。')
    
    doc.add_heading('Implementation Details', level=1)
    doc.add_paragraph('我们将致力于为客户提供无可比拟的体验。')
    
    doc.save(path)
    print(f"Created integrated test source: {path}")

def run_integrated_test():
    source_file = "test_sprint2_source.docx"
    create_integrated_test_source(source_file)
    
    # 1. Semantic Chunking
    print("\n--- Phase 1: Semantic Chunking ---")
    chunker = DynamicSemanticChunker(embedder=MockEmbedder(), threshold=0.3, target_dir="data/test_chunks")
    chunks = chunker.chunk_document(source_file)
    
    # 2. Tone Calibration
    print("\n--- Phase 2: Tone Calibration ---")
    calibrator = ToneCalibrator()
    
    for c in chunks:
        print(f"\nProcessing Chunk {c['id']} from {c['path']}...")
        original_text = c['text']
        calibrated_text = calibrator.calibrate_text(original_text)
        
        print(f"Original Text Snip: {original_text[:50]}...")
        print(f"Calibrated Text Snip: {calibrated_text[:50]}...")
        
        # Verify specific replacements
        if "综上所述" in original_text and "综上所述" not in calibrated_text:
             print("CHECK: Forbidden word removed.")
        if "毫秒级响应" in calibrated_text:
             print("CHECK: Fuzzy word '很快' replaced by '毫秒级响应'.")

    print("\n--- Integrated Test Completed ---")

if __name__ == "__main__":
    run_integrated_test()
